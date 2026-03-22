"""Export utilities for Word, JSON, and structured Markdown outputs."""

from __future__ import annotations

import json
from datetime import datetime
from pathlib import Path

from utils.config import AppConfig
from utils.model_types import WorkflowResult
from utils.validators import sanitize_filename


class ExportService:
    """Handles generation of downloadable artifacts."""

    def __init__(self, config: AppConfig, logger) -> None:
        self.config = config
        self.logger = logger

    def export_json(self, result: WorkflowResult, source_name: str) -> Path:
        path = self.config.output_dir / f"{self._stem(source_name)}_workflow.json"
        with path.open("w", encoding="utf-8") as handle:
            json.dump(result.structured_json, handle, indent=2, ensure_ascii=False)
        self.logger.info("JSON export generated: %s", path)
        return path

    def export_markdown(self, result: WorkflowResult, source_name: str) -> Path:
        path = self.config.output_dir / f"{self._stem(source_name)}_sop.md"
        with path.open("w", encoding="utf-8") as handle:
            handle.write(result.sop_markdown)
        self.logger.info("Markdown export generated: %s", path)
        return path

    def export_word(self, result: WorkflowResult, source_name: str) -> Path:
        """Create a DOCX export with fallback to TXT if python-docx is unavailable."""
        stem = self._stem(source_name)
        docx_path = self.config.output_dir / f"{stem}_workflow.docx"

        try:
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            from docx.shared import Pt

            document = Document()
            title = document.add_paragraph()
            title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = title.add_run("Canberbyte Technologies - Workflow Export")
            run.bold = True
            run.font.size = Pt(18)

            subtitle = document.add_paragraph(
                f"Source: {source_name}\nGenerated: {datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')}"
            )
            subtitle.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            document.add_heading("Standard Operating Procedure", level=1)
            for line in result.sop_markdown.splitlines():
                if line.strip():
                    document.add_paragraph(line)

            document.add_heading("Structured Workflow Steps", level=1)
            table = document.add_table(rows=1, cols=4)
            hdr = table.rows[0].cells
            hdr[0].text = "Step ID"
            hdr[1].text = "Title"
            hdr[2].text = "Role"
            hdr[3].text = "Description"

            for step in result.workflow_steps:
                row = table.add_row().cells
                row[0].text = step.step_id
                row[1].text = step.title
                row[2].text = step.role
                row[3].text = step.description

            document.save(docx_path)
            self.logger.info("Word export generated: %s", docx_path)
            return docx_path

        except Exception as exc:
            self.logger.warning("DOCX export failed; writing fallback TXT. Reason: %s", exc)
            txt_path = self.config.output_dir / f"{stem}_workflow.txt"
            with txt_path.open("w", encoding="utf-8") as handle:
                handle.write(result.sop_markdown)
            return txt_path

    @staticmethod
    def _stem(source_name: str) -> str:
        safe = sanitize_filename(Path(source_name).stem)
        return safe.lower()
